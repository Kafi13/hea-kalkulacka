"""
HEA Expert: Hydrogen & Thermodynamics Calculator
Streamlit aplikace pro výpočet termodynamických parametrů vysokoentropických slitin
a predikci vodíkové afinity.

Autor: RTI ZČU
Verze: 1.0
"""

import streamlit as st
import numpy as np
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
from plotly.subplots import make_subplots
import re
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from datetime import datetime

# ============== KONFIGURACE STRÁNKY ==============
st.set_page_config(
    page_title="HEA Expert Calculator",
    page_icon="🧬",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ============== CUSTOM CSS ==============
st.markdown("""
<style>
    .main-header {
        font-size: 2.5rem;
        font-weight: 700;
        background: linear-gradient(90deg, #E91E63, #2196F3);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        text-align: center;
        margin-bottom: 0.5rem;
    }
    .sub-header {
        text-align: center;
        color: #888;
        margin-bottom: 2rem;
    }
    .metric-card {
        background: linear-gradient(135deg, rgba(255,255,255,0.1), rgba(255,255,255,0.05));
        border-radius: 10px;
        padding: 1rem;
        border: 1px solid rgba(255,255,255,0.1);
    }
    .stMetric {
        background: rgba(255,255,255,0.05);
        padding: 1rem;
        border-radius: 10px;
    }
</style>
""", unsafe_allow_html=True)

# ============== DATABÁZE PRVKŮ ==============
ELEMENTS = {
    "Al": {"name": "Hliník", "group": "Lehké kovy", "r": 143, "VEC": 3, "Tm": 933, "phi": 4.20, "nWS": 1.39, "V23": 4.64, "deltaH_H": -4.0, "deltaH_inf": 6.0, "color": "#90CAF9"},
    "Mg": {"name": "Hořčík", "group": "Lehké kovy", "r": 160, "VEC": 2, "Tm": 923, "phi": 3.45, "nWS": 1.17, "V23": 5.89, "deltaH_H": -37.0, "deltaH_inf": -19.0, "color": "#81D4FA"},
    "Si": {"name": "Křemík", "group": "Polokovy", "r": 118, "VEC": 4, "Tm": 1687, "phi": 4.70, "nWS": 1.50, "V23": 4.20, "deltaH_H": 0.0, "deltaH_inf": 46.0, "color": "#B39DDB"},
    "Sc": {"name": "Skandium", "group": "3d přechodné", "r": 162, "VEC": 3, "Tm": 1814, "phi": 3.25, "nWS": 1.27, "V23": 6.09, "deltaH_H": -100.0, "deltaH_inf": -78.0, "color": "#CE93D8"},
    "Ti": {"name": "Titan", "group": "3d přechodné", "r": 147, "VEC": 4, "Tm": 1941, "phi": 3.80, "nWS": 1.52, "V23": 4.12, "deltaH_H": -67.0, "deltaH_inf": -52.0, "color": "#7986CB"},
    "V": {"name": "Vanad", "group": "3d přechodné", "r": 134, "VEC": 5, "Tm": 2183, "phi": 4.25, "nWS": 1.64, "V23": 3.71, "deltaH_H": -34.0, "deltaH_inf": -26.0, "color": "#64B5F6"},
    "Cr": {"name": "Chrom", "group": "3d přechodné", "r": 128, "VEC": 6, "Tm": 2180, "phi": 4.65, "nWS": 1.73, "V23": 3.35, "deltaH_H": 6.0, "deltaH_inf": 25.0, "color": "#4FC3F7"},
    "Mn": {"name": "Mangan", "group": "3d přechodné", "r": 127, "VEC": 7, "Tm": 1519, "phi": 4.45, "nWS": 1.61, "V23": 3.37, "deltaH_H": -8.0, "deltaH_inf": 8.0, "color": "#4DD0E1"},
    "Fe": {"name": "Železo", "group": "3d přechodné", "r": 126, "VEC": 8, "Tm": 1811, "phi": 4.93, "nWS": 1.77, "V23": 3.29, "deltaH_H": 14.0, "deltaH_inf": 29.0, "color": "#4DB6AC"},
    "Co": {"name": "Kobalt", "group": "3d přechodné", "r": 125, "VEC": 9, "Tm": 1768, "phi": 5.10, "nWS": 1.79, "V23": 3.18, "deltaH_H": 11.0, "deltaH_inf": 25.0, "color": "#81C784"},
    "Ni": {"name": "Nikl", "group": "3d přechodné", "r": 124, "VEC": 10, "Tm": 1728, "phi": 5.20, "nWS": 1.80, "V23": 3.09, "deltaH_H": -2.0, "deltaH_inf": 12.0, "color": "#AED581"},
    "Cu": {"name": "Měď", "group": "3d přechodné", "r": 128, "VEC": 11, "Tm": 1358, "phi": 4.55, "nWS": 1.47, "V23": 3.37, "deltaH_H": 21.0, "deltaH_inf": 46.0, "color": "#FFB74D"},
    "Zn": {"name": "Zinek", "group": "3d přechodné", "r": 134, "VEC": 12, "Tm": 693, "phi": 4.10, "nWS": 1.32, "V23": 4.02, "deltaH_H": 13.0, "deltaH_inf": 35.0, "color": "#FF8A65"},
    "Y": {"name": "Yttrium", "group": "4d přechodné", "r": 180, "VEC": 3, "Tm": 1799, "phi": 3.20, "nWS": 1.21, "V23": 8.26, "deltaH_H": -114.0, "deltaH_inf": -91.0, "color": "#F48FB1"},
    "Zr": {"name": "Zirkonium", "group": "4d přechodné", "r": 160, "VEC": 4, "Tm": 2128, "phi": 3.45, "nWS": 1.41, "V23": 5.89, "deltaH_H": -80.0, "deltaH_inf": -65.0, "color": "#CE93D8"},
    "Nb": {"name": "Niob", "group": "4d přechodné", "r": 146, "VEC": 5, "Tm": 2750, "phi": 4.05, "nWS": 1.62, "V23": 4.93, "deltaH_H": -40.0, "deltaH_inf": -32.0, "color": "#B39DDB"},
    "Mo": {"name": "Molybden", "group": "4d přechodné", "r": 139, "VEC": 6, "Tm": 2896, "phi": 4.65, "nWS": 1.77, "V23": 4.24, "deltaH_H": 20.0, "deltaH_inf": 37.0, "color": "#9FA8DA"},
    "Pd": {"name": "Palladium", "group": "4d přechodné", "r": 137, "VEC": 10, "Tm": 1828, "phi": 5.45, "nWS": 1.67, "V23": 4.03, "deltaH_H": -10.0, "deltaH_inf": 1.0, "color": "#80DEEA"},
    "Ag": {"name": "Stříbro", "group": "4d přechodné", "r": 144, "VEC": 11, "Tm": 1235, "phi": 4.45, "nWS": 1.39, "V23": 4.61, "deltaH_H": 31.0, "deltaH_inf": 62.0, "color": "#B0BEC5"},
    "Hf": {"name": "Hafnium", "group": "5d přechodné", "r": 159, "VEC": 4, "Tm": 2506, "phi": 3.55, "nWS": 1.46, "V23": 5.76, "deltaH_H": -65.0, "deltaH_inf": -52.0, "color": "#FFCC80"},
    "Ta": {"name": "Tantal", "group": "5d přechodné", "r": 146, "VEC": 5, "Tm": 3290, "phi": 4.15, "nWS": 1.63, "V23": 4.93, "deltaH_H": -38.0, "deltaH_inf": -30.0, "color": "#FFAB91"},
    "W": {"name": "Wolfram", "group": "5d přechodné", "r": 139, "VEC": 6, "Tm": 3695, "phi": 4.80, "nWS": 1.81, "V23": 4.24, "deltaH_H": 40.0, "deltaH_inf": 56.0, "color": "#BCAAA4"},
    "Re": {"name": "Rhenium", "group": "5d přechodné", "r": 137, "VEC": 7, "Tm": 3459, "phi": 5.20, "nWS": 1.86, "V23": 4.03, "deltaH_H": 25.0, "deltaH_inf": 40.0, "color": "#B0BEC5"},
    "Pt": {"name": "Platina", "group": "5d přechodné", "r": 139, "VEC": 10, "Tm": 2041, "phi": 5.65, "nWS": 1.78, "V23": 4.24, "deltaH_H": -5.0, "deltaH_inf": 10.0, "color": "#E0E0E0"},
    "Au": {"name": "Zlato", "group": "5d přechodné", "r": 144, "VEC": 11, "Tm": 1337, "phi": 5.15, "nWS": 1.57, "V23": 4.61, "deltaH_H": 25.0, "deltaH_inf": 52.0, "color": "#FFD54F"},
    "La": {"name": "Lanthan", "group": "Lanthanoidy", "r": 187, "VEC": 3, "Tm": 1193, "phi": 3.05, "nWS": 1.18, "V23": 9.04, "deltaH_H": -100.0, "deltaH_inf": -80.0, "color": "#FFE082"},
    "Ce": {"name": "Cer", "group": "Lanthanoidy", "r": 182, "VEC": 3, "Tm": 1068, "phi": 3.18, "nWS": 1.19, "V23": 8.54, "deltaH_H": -95.0, "deltaH_inf": -75.0, "color": "#FFF59D"},
}

# ============== BINÁRNÍ ENTALPIE ==============
BINARY_ENTHALPIES = {
    "Ti-V": -2, "Ti-Cr": -7, "Ti-Mn": -8, "Ti-Fe": -17, "Ti-Co": -28,
    "Ti-Ni": -35, "Ti-Cu": -9, "Ti-Zn": -15, "Ti-Al": -30, "Ti-Zr": 0,
    "Ti-Nb": 2, "Ti-Mo": -4, "Ti-Hf": 0, "Ti-Ta": 1, "Ti-W": -6,
    "Ti-La": 13, "Ti-Ce": 10, "Ti-Y": 8, "Ti-Sc": 4, "Ti-Pd": -63,
    "Ti-Mg": 16, "Ti-Si": -66,
    "V-Cr": -2, "V-Mn": -2, "V-Fe": -7, "V-Co": -14, "V-Ni": -18,
    "V-Cu": 5, "V-Zn": -4, "V-Al": -16, "V-Zr": -4, "V-Nb": -1,
    "V-Mo": 0, "V-Hf": -2, "V-Ta": -1, "V-W": -1, "V-La": 22,
    "V-Pd": -44, "V-Si": -45,
    "Cr-Mn": 2, "Cr-Fe": -1, "Cr-Co": -4, "Cr-Ni": -7, "Cr-Cu": 12,
    "Cr-Zn": 5, "Cr-Al": -10, "Cr-Zr": -12, "Cr-Nb": -7, "Cr-Mo": 0,
    "Cr-Hf": -9, "Cr-Ta": -7, "Cr-W": 1, "Cr-La": 27, "Cr-Pd": -27,
    "Cr-Si": -37,
    "Mn-Fe": 0, "Mn-Co": -5, "Mn-Ni": -8, "Mn-Cu": 4, "Mn-Zn": -4,
    "Mn-Al": -19, "Mn-Zr": -20, "Mn-Nb": -13, "Mn-Mo": -5,
    "Mn-Pd": -23, "Mn-Si": -45,
    "Fe-Co": -1, "Fe-Ni": -2, "Fe-Cu": 13, "Fe-Zn": -1, "Fe-Al": -11,
    "Fe-Zr": -25, "Fe-Nb": -16, "Fe-Mo": -2, "Fe-Hf": -21, "Fe-Ta": -15,
    "Fe-W": 0, "Fe-La": 18, "Fe-Pd": -4, "Fe-Si": -35,
    "Co-Ni": 0, "Co-Cu": 6, "Co-Zn": -7, "Co-Al": -19, "Co-Zr": -41,
    "Co-Nb": -25, "Co-Mo": -5, "Co-Hf": -35, "Co-Ta": -24, "Co-W": -1,
    "Co-La": 7, "Co-Pd": 0, "Co-Si": -38,
    "Ni-Cu": 4, "Ni-Zn": -9, "Ni-Al": -22, "Ni-Zr": -49, "Ni-Nb": -30,
    "Ni-Mo": -7, "Ni-Hf": -42, "Ni-Ta": -29, "Ni-W": -3, "Ni-La": -4,
    "Ni-Pd": 0, "Ni-Si": -40, "Ni-Mg": -4,
    "Cu-Zn": -6, "Cu-Al": -1, "Cu-Zr": -23, "Cu-Nb": -3, "Cu-Mo": 19,
    "Cu-Hf": -21, "Cu-Ta": -2, "Cu-W": 22, "Cu-La": -15, "Cu-Pd": -14,
    "Cu-Mg": -3, "Cu-Si": -10,
    "Zn-Al": 1, "Zn-Zr": -37, "Zn-Nb": -19, "Zn-La": -29, "Zn-Mg": 4,
    "Al-Zr": -44, "Al-Nb": -18, "Al-Mo": -5, "Al-Hf": -39, "Al-Ta": -19,
    "Al-W": -2, "Al-La": -38, "Al-Ce": -38, "Al-Y": -38, "Al-Sc": -38,
    "Al-Pd": -55, "Al-Mg": 2, "Al-Si": -4,
    "Zr-Nb": 4, "Zr-Mo": -6, "Zr-Hf": 0, "Zr-Ta": 3, "Zr-W": -9,
    "Zr-La": 10, "Zr-Pd": -91, "Zr-Si": -84,
    "Nb-Mo": -6, "Nb-Hf": 4, "Nb-Ta": 0, "Nb-W": -8, "Nb-La": 16,
    "Nb-Pd": -64, "Nb-Si": -56,
    "Mo-Hf": -4, "Mo-Ta": -5, "Mo-W": 0, "Mo-La": 32, "Mo-Pd": -18,
    "Hf-Ta": 3, "Hf-W": -6, "Hf-La": 11, "Hf-Pd": -84,
    "Ta-W": -7, "Ta-La": 17, "Ta-Pd": -60,
    "W-La": 36, "W-Pd": -12,
    "La-Ce": 0, "La-Pd": -76, "La-Mg": 6,
    "Ce-Pd": -75,
    "Pd-Si": -55, "Pd-Mg": -45,
    "Y-Zr": 9, "Y-Sc": 0,
    "Sc-Zr": 5,
    "Mg-Si": -3
}


def get_binary_enthalpy(el1: str, el2: str) -> float:
    """Získání binární entalpie míšení pro pár prvků."""
    key1 = f"{el1}-{el2}"
    key2 = f"{el2}-{el1}"
    
    if key1 in BINARY_ENTHALPIES:
        return BINARY_ENTHALPIES[key1]
    if key2 in BINARY_ENTHALPIES:
        return BINARY_ENTHALPIES[key2]
    
    # Miedema odhad
    a = ELEMENTS.get(el1)
    b = ELEMENTS.get(el2)
    if a and b and all(k in a for k in ["phi", "nWS", "V23"]):
        P, Q = 14.2, 9.4
        d_phi = a["phi"] - b["phi"]
        d_n = a["nWS"] - b["nWS"]
        V_avg = 2 * a["V23"] * b["V23"] / (a["V23"] + b["V23"])
        return round(V_avg * (-P * d_phi**2 + Q * d_n**2))
    
    return 0


def parse_formula(formula: str) -> dict:
    """
    Parser chemických vzorců HEA.
    Podporované formáty:
    - TiVCrNb (ekvlatomární)
    - Ti0.2V0.2Cr0.2Ni0.2Nb0.2 (atomové zlomky)
    - Ti20V20Cr20Ni20Nb20 (procenta)
    - (TiV)95Ni5 (skupina s procentem)
    - Ti2V3Cr1 (poměry)
    """
    formula = formula.strip().replace(" ", "")
    if not formula:
        return None
    
    result = {}
    remaining = formula
    
    # Regex pro skupiny v závorkách
    group_pattern = re.compile(r'^\(([^)]+)\)(\d*\.?\d*)')
    # Regex pro prvky
    element_pattern = re.compile(r'^([A-Z][a-z]?)(\d*\.?\d*)')
    
    while remaining:
        # Zkus najít skupinu v závorkách
        match = group_pattern.match(remaining)
        if match:
            group_content = match.group(1)
            multiplier = float(match.group(2)) if match.group(2) else 1.0
            
            # Parsuj prvky ve skupině
            group_elements = {}
            group_remaining = group_content
            while group_remaining:
                el_match = element_pattern.match(group_remaining)
                if el_match and el_match.group(1) in ELEMENTS:
                    el = el_match.group(1)
                    amount = float(el_match.group(2)) if el_match.group(2) else 1.0
                    group_elements[el] = group_elements.get(el, 0) + amount
                    group_remaining = group_remaining[len(el_match.group(0)):]
                else:
                    group_remaining = group_remaining[1:]
            
            for el, amount in group_elements.items():
                result[el] = result.get(el, 0) + amount * multiplier
            
            remaining = remaining[len(match.group(0)):]
            continue
        
        # Zkus najít jednotlivý prvek
        match = element_pattern.match(remaining)
        if match and match.group(1) in ELEMENTS:
            el = match.group(1)
            amount = float(match.group(2)) if match.group(2) else 1.0
            result[el] = result.get(el, 0) + amount
            remaining = remaining[len(match.group(0)):]
            continue
        
        # Neplatný znak - přeskoč
        remaining = remaining[1:]
    
    if not result:
        return None
    
    # Normalizuj na procenta
    total = sum(result.values())
    return {el: (amount / total) * 100 for el, amount in result.items()}


def calculate_thermodynamics(elements: list, compositions: dict) -> dict:
    """Výpočet všech termodynamických parametrů."""
    n = len(elements)
    if n < 2:
        return None
    
    # Normalizace
    total = sum(compositions.get(el, 0) for el in elements)
    if total == 0:
        return None
    norm_comp = {el: compositions.get(el, 0) / total for el in elements}
    
    R = 8.314
    
    # Průměrné vlastnosti
    r_avg = sum(norm_comp[el] * ELEMENTS[el]["r"] for el in elements)
    Tm_avg = sum(norm_comp[el] * ELEMENTS[el]["Tm"] for el in elements)
    VEC = sum(norm_comp[el] * ELEMENTS[el]["VEC"] for el in elements)
    deltaH_inf = sum(norm_comp[el] * ELEMENTS[el]["deltaH_inf"] for el in elements)
    deltaH_f = sum(norm_comp[el] * ELEMENTS[el]["deltaH_H"] for el in elements)
    
    # Entropie míšení
    S_mix = 0
    for el in elements:
        c = norm_comp[el]
        if c > 0:
            S_mix -= c * np.log(c)
    S_mix *= R
    
    # Entalpie míšení (Miedema)
    H_mix = 0
    for i in range(n):
        for j in range(i + 1, n):
            el1, el2 = elements[i], elements[j]
            c1, c2 = norm_comp[el1], norm_comp[el2]
            Hij = get_binary_enthalpy(el1, el2)
            H_mix += 4 * Hij * c1 * c2
    
    # Atomový mismatch δ
    delta = 0
    for el in elements:
        c = norm_comp[el]
        ri = ELEMENTS[el]["r"]
        delta += c * (1 - ri / r_avg) ** 2
    delta = 100 * np.sqrt(delta)
    
    # Omega parametr
    Omega = (Tm_avg * S_mix / 1000) / abs(H_mix) if abs(H_mix) > 0.01 else float('inf')
    
    return {
        "S_mix": S_mix / R,  # v jednotkách R
        "H_mix": H_mix,
        "delta": delta,
        "Omega": Omega,
        "VEC": VEC,
        "r_avg": r_avg,
        "Tm_avg": Tm_avg,
        "deltaH_inf": deltaH_inf,
        "deltaH_f": deltaH_f,
        "norm_comp": norm_comp
    }


def predict_structure(results: dict) -> dict:
    """Predikce krystalové struktury."""
    delta = results["delta"]
    Omega = results["Omega"]
    VEC = results["VEC"]
    
    if delta > 6.6:
        return {"structure": "Amorfní/Intermetalická", "confidence": "vysoká", "icon": "⚠️", "color": "#E53935"}
    if Omega < 1.1:
        return {"structure": "Intermetalické sloučeniny", "confidence": "střední", "icon": "⚠️", "color": "#FB8C00"}
    if VEC < 6.87:
        return {"structure": "BCC (kubická prostorově centrovaná)", "confidence": "vysoká", "icon": "🔷", "color": "#2196F3"}
    if VEC > 8.0:
        return {"structure": "FCC (kubická plošně centrovaná)", "confidence": "vysoká", "icon": "🔶", "color": "#FF9800"}
    return {"structure": "BCC + FCC (duální fáze)", "confidence": "střední", "icon": "🔷🔶", "color": "#9C27B0"}


def classify_hydrogen(results: dict) -> dict:
    """Klasifikace vodíkové afinity."""
    deltaH_f = results["deltaH_f"]
    
    if deltaH_f < -60:
        return {
            "category": "Vodíková past",
            "color": "#E53935",
            "icon": "🔥",
            "description": "Příliš stabilní hydrid. Desorpce vyžaduje >300°C. Nevhodné pro reverzibilní skladování."
        }
    if deltaH_f < -20:
        return {
            "category": "Skladovací materiál",
            "color": "#43A047",
            "icon": "✓",
            "description": "Ideální rozsah pro reverzibilní skladování H₂. Absorpce/desorpce při provozních teplotách."
        }
    if deltaH_f < 0:
        return {
            "category": "Mírná absorpce",
            "color": "#FB8C00",
            "icon": "~",
            "description": "Slabá tvorba hydridů. Možné použití jako membrána nebo ochranný povlak."
        }
    return {
        "category": "Vodíková bariéra",
        "color": "#1E88E5",
        "icon": "🛡️",
        "description": "Materiál odolný vůči absorpci vodíku. Vhodný pro aplikace vyžadující odolnost vůči vodíkové křehkosti."
    }


def generate_formula_string(elements: list, norm_comp: dict) -> str:
    """Generování vzorce ze složení."""
    parts = []
    for el in elements:
        pct = norm_comp[el] * 100
        if pct > 0:
            if abs(pct - round(pct)) < 0.1 and round(pct) == 100 / len(elements):
                parts.append(el)
            else:
                parts.append(f"{el}{pct:.1f}")
    return "".join(parts) if parts else "".join(elements)


def create_word_report(elements: list, results: dict, structure: dict, hydrogen: dict) -> BytesIO:
    """Vytvoření Word reportu."""
    doc = Document()
    
    # Nastavení stylů
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    
    formula = generate_formula_string(elements, results["norm_comp"])
    today = datetime.now().strftime("%d.%m.%Y")
    
    # Titulní strana
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("TERMODYNAMICKÁ ANALÝZA\n")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(21, 101, 192)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("VYSOKOENTROPICKÉ SLITINY\n& PREDIKCE VODÍKOVÉ AFINITY")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(25, 118, 210)
    
    doc.add_paragraph()
    
    formula_p = doc.add_paragraph()
    formula_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = formula_p.add_run(formula)
    run.bold = True
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(21, 101, 192)
    
    doc.add_paragraph()
    
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run(f"Datum analýzy: {today}")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(100, 100, 100)
    
    doc.add_page_break()
    
    # 1. Složení slitiny
    doc.add_heading("1. Složení slitiny", level=1)
    doc.add_paragraph("Analyzovaná vysokoentropická slitina (HEA) má následující chemické složení:")
    
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    headers = ["Prvek", "Symbol", "Složení [at.%]", "r [pm]", "VEC", "Tₘ [K]"]
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
    
    for el in elements:
        row = table.add_row().cells
        row[0].text = ELEMENTS[el]["name"]
        row[1].text = el
        row[2].text = f"{results['norm_comp'][el] * 100:.2f}"
        row[3].text = str(ELEMENTS[el]["r"])
        row[4].text = str(ELEMENTS[el]["VEC"])
        row[5].text = str(ELEMENTS[el]["Tm"])
    
    doc.add_paragraph()
    
    # 2. Termodynamické parametry
    doc.add_heading("2. Termodynamické parametry", level=1)
    doc.add_paragraph("Výpočet termodynamických parametrů podle empirických pravidel:")
    
    params_table = doc.add_table(rows=6, cols=4)
    params_table.style = 'Table Grid'
    params_data = [
        ["Parametr", "Hodnota", "Kritérium", "Hodnocení"],
        ["ΔSₘᵢₓ [R]", f"{results['S_mix']:.3f}", "> 1.5 R", "✓" if results['S_mix'] > 1.5 else "✗"],
        ["ΔHₘᵢₓ [kJ/mol]", f"{results['H_mix']:.2f}", "-11.6 až 3.2", "✓" if -11.6 <= results['H_mix'] <= 3.2 else "✗"],
        ["δ [%]", f"{results['delta']:.3f}", "< 6.6%", "✓" if results['delta'] < 6.6 else "✗"],
        ["Ω", f"{results['Omega']:.3f}" if results['Omega'] != float('inf') else "∞", "> 1.1", "✓" if results['Omega'] > 1.1 else "✗"],
        ["VEC", f"{results['VEC']:.3f}", "< 6.87 → BCC", "BCC" if results['VEC'] < 6.87 else ("FCC" if results['VEC'] > 8 else "BCC+FCC")]
    ]
    for i, row_data in enumerate(params_data):
        row = params_table.rows[i].cells
        for j, cell_data in enumerate(row_data):
            row[j].text = cell_data
            if i == 0:
                row[j].paragraphs[0].runs[0].bold = True
    
    doc.add_paragraph()
    doc.add_paragraph(f"Průměrný atomový poloměr: r̄ = {results['r_avg']:.2f} pm")
    doc.add_paragraph(f"Průměrná teplota tání: Tₘ = {results['Tm_avg']:.0f} K ({results['Tm_avg'] - 273.15:.0f} °C)")
    
    # 3. Predikce struktury
    doc.add_heading("3. Predikce krystalové struktury", level=1)
    p = doc.add_paragraph()
    run = p.add_run(f"{structure['icon']} Predikovaná struktura: {structure['structure']}")
    run.bold = True
    run.font.size = Pt(14)
    doc.add_paragraph(f"Spolehlivost predikce: {structure['confidence']}")
    
    doc.add_page_break()
    
    # 4. Vodíková afinita
    doc.add_heading("4. Analýza vodíkové afinity", level=1)
    doc.add_paragraph("Predikce interakce slitiny s vodíkem podle modelu Griessen-Driessen:")
    
    h_table = doc.add_table(rows=3, cols=3)
    h_table.style = 'Table Grid'
    h_data = [
        ["Parametr", "Hodnota", "Popis"],
        ["ΔH∞ [kJ/mol H]", f"{results['deltaH_inf']:.2f}", "Entalpie roztoku H při nekonečném zředění"],
        ["ΔHf [kJ/mol H]", f"{results['deltaH_f']:.2f}", "Entalpie tvorby hydridu"]
    ]
    for i, row_data in enumerate(h_data):
        row = h_table.rows[i].cells
        for j, cell_data in enumerate(row_data):
            row[j].text = cell_data
            if i == 0:
                row[j].paragraphs[0].runs[0].bold = True
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(f"{hydrogen['icon']} Klasifikace: {hydrogen['category']}")
    run.bold = True
    run.font.size = Pt(14)
    doc.add_paragraph(hydrogen['description'])
    
    # 5. Závěr
    doc.add_heading("5. Závěr", level=1)
    doc.add_paragraph(f"Analyzovaná slitina {formula} vykazuje následující charakteristiky:")
    doc.add_paragraph(f"• Predikovaná krystalová struktura: {structure['structure']}")
    doc.add_paragraph(f"• Vodíková afinita: {hydrogen['category']} (ΔHf = {results['deltaH_f']:.1f} kJ/mol H)")
    doc.add_paragraph(f"• Splnění kritérií pro tuhý roztok: δ < 6.6% ({'ANO' if results['delta'] < 6.6 else 'NE'}), Ω > 1.1 ({'ANO' if results['Omega'] > 1.1 else 'NE'})")
    
    # Reference
    doc.add_heading("Reference", level=1)
    doc.add_paragraph("[1] Zhang, Y. et al. (2014). Microstructures and properties of high-entropy alloys. Progress in Materials Science, 61, 1-93.")
    doc.add_paragraph("[2] Griessen, R., & Driessen, A. (1984). Heat of formation and band structure of binary and ternary metal hydrides. Physical Review B, 30(8), 4372.")
    doc.add_paragraph("[3] Miedema, A. R. (1973). A simple model for alloys. Philips Technical Review, 33, 149-160.")
    
    # Uložení do BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ============== HLAVNÍ APLIKACE ==============
def main():
    # Header
    st.markdown('<h1 class="main-header">🧬 HEA Expert: Hydrogen & Thermodynamics</h1>', unsafe_allow_html=True)
    st.markdown('<p class="sub-header">Výpočet termodynamických parametrů vysokoentropických slitin a predikce vodíkové afinity</p>', unsafe_allow_html=True)
    
    # Inicializace session state
    if 'elements' not in st.session_state:
        st.session_state.elements = ['Ti', 'V', 'Cr', 'Ni', 'Nb']
    if 'compositions' not in st.session_state:
        st.session_state.compositions = {el: 20.0 for el in st.session_state.elements}
    
    # Vstupní pole pro vzorec
    col1, col2, col3 = st.columns([3, 1, 1])
    with col1:
        formula_input = st.text_input(
            "Zadejte složení (např. TiVCrNb, (TiV)95Ni5, Ti0.2V0.2Cr0.2Ni0.2Nb0.2):",
            value="TiVCrNiNb",
            key="formula_input"
        )
    with col2:
        if st.button("🚀 Analyzovat", type="primary", use_container_width=True):
            parsed = parse_formula(formula_input)
            if parsed and len(parsed) >= 2:
                st.session_state.elements = list(parsed.keys())
                st.session_state.compositions = parsed
                st.rerun()
            else:
                st.error("Neplatný vzorec nebo méně než 2 prvky!")
    
    st.divider()
    
    # Hlavní layout
    col_left, col_center, col_right = st.columns([1.2, 1.5, 1.3])
    
    # LEVÝ PANEL - Výběr prvků
    with col_left:
        st.subheader("🔬 Výběr prvků")
        
        # Multiselect pro prvky
        available_elements = list(ELEMENTS.keys())
        selected = st.multiselect(
            "Vyberte prvky (2-8):",
            options=available_elements,
            default=st.session_state.elements,
            max_selections=8,
            format_func=lambda x: f"{x} ({ELEMENTS[x]['name']})"
        )
        
        if len(selected) >= 2:
            st.session_state.elements = selected
            # Aktualizuj kompozice pro nové prvky
            for el in selected:
                if el not in st.session_state.compositions:
                    st.session_state.compositions[el] = 100 / len(selected)
            # Odstraň staré prvky
            st.session_state.compositions = {k: v for k, v in st.session_state.compositions.items() if k in selected}
        
        st.divider()
        
        # Tlačítko pro ekvlatomární složení
        if st.button("⚖️ Nastavit ekvlatomární", use_container_width=True):
            eq_val = 100 / len(st.session_state.elements)
            st.session_state.compositions = {el: eq_val for el in st.session_state.elements}
            st.rerun()
        
        st.divider()
        
        # Slidery pro složení
        st.markdown("**Složení [at.%]:**")
        for el in st.session_state.elements:
            new_val = st.slider(
                f"{el} ({ELEMENTS[el]['name']})",
                min_value=0.0,
                max_value=100.0,
                value=st.session_state.compositions.get(el, 20.0),
                step=0.5,
                key=f"slider_{el}"
            )
            st.session_state.compositions[el] = new_val
        
        # Pie chart
        if st.session_state.elements:
            total = sum(st.session_state.compositions.get(el, 0) for el in st.session_state.elements)
            if total > 0:
                pie_data = pd.DataFrame({
                    'Prvek': st.session_state.elements,
                    'Podíl': [st.session_state.compositions.get(el, 0) / total * 100 for el in st.session_state.elements],
                    'Barva': [ELEMENTS[el]['color'] for el in st.session_state.elements]
                })
                
                fig_pie = px.pie(
                    pie_data,
                    values='Podíl',
                    names='Prvek',
                    color='Prvek',
                    color_discrete_map={el: ELEMENTS[el]['color'] for el in st.session_state.elements},
                    hole=0.4
                )
                fig_pie.update_traces(textposition='outside', textinfo='label+percent')
                fig_pie.update_layout(
                    showlegend=False,
                    margin=dict(t=20, b=20, l=20, r=20),
                    height=300
                )
                st.plotly_chart(fig_pie, use_container_width=True)
    
    # Výpočet výsledků
    results = calculate_thermodynamics(st.session_state.elements, st.session_state.compositions)
    
    if results:
        structure = predict_structure(results)
        hydrogen = classify_hydrogen(results)
        formula_str = generate_formula_string(st.session_state.elements, results["norm_comp"])
        
        # STŘEDNÍ PANEL - Výsledky
        with col_center:
            st.subheader("📊 Termodynamické parametry")
            
            # Metriky v gridu
            m1, m2, m3 = st.columns(3)
            with m1:
                st.metric("ΔSₘᵢₓ [R]", f"{results['S_mix']:.3f}", 
                         delta="✓" if results['S_mix'] > 1.5 else "⚠")
            with m2:
                st.metric("ΔHₘᵢₓ [kJ/mol]", f"{results['H_mix']:.2f}",
                         delta="✓" if -11.6 <= results['H_mix'] <= 3.2 else "⚠")
            with m3:
                st.metric("δ [%]", f"{results['delta']:.3f}",
                         delta="✓" if results['delta'] < 6.6 else "⚠")
            
            m4, m5, m6 = st.columns(3)
            with m4:
                omega_str = f"{results['Omega']:.3f}" if results['Omega'] != float('inf') else "∞"
                st.metric("Ω", omega_str,
                         delta="✓" if results['Omega'] > 1.1 else "⚠")
            with m5:
                vec_phase = "BCC" if results['VEC'] < 6.87 else ("FCC" if results['VEC'] > 8 else "Dual")
                st.metric("VEC", f"{results['VEC']:.3f}", delta=vec_phase)
            with m6:
                st.metric("Tₘ [K]", f"{results['Tm_avg']:.0f}")
            
            st.divider()
            
            # Predikce struktury
            st.markdown(f"### {structure['icon']} Predikovaná struktura")
            st.info(f"**{structure['structure']}**\n\nSpolehlivost: {structure['confidence']}")
            
            st.divider()
            
            # Diagram stability
            st.markdown("### 📈 Diagram stability")
            
            fig_scatter = go.Figure()
            
            # Referenční oblasti
            fig_scatter.add_shape(type="rect", x0=-11.6, x1=3.2, y0=0, y1=6.6,
                                 fillcolor="rgba(76, 175, 80, 0.1)", line=dict(width=0))
            fig_scatter.add_hline(y=6.6, line_dash="dash", line_color="red",
                                 annotation_text="δ = 6.6%")
            fig_scatter.add_vline(x=-11.6, line_dash="dash", line_color="blue")
            fig_scatter.add_vline(x=3.2, line_dash="dash", line_color="blue")
            
            # Bod slitiny
            fig_scatter.add_trace(go.Scatter(
                x=[results['H_mix']],
                y=[results['delta']],
                mode='markers',
                marker=dict(size=20, color='#4FC3F7', symbol='diamond'),
                name=formula_str,
                text=[f"{formula_str}<br>ΔHmix={results['H_mix']:.2f}<br>δ={results['delta']:.3f}%"],
                hoverinfo='text'
            ))
            
            fig_scatter.update_layout(
                xaxis_title="ΔHₘᵢₓ [kJ/mol]",
                yaxis_title="δ [%]",
                xaxis=dict(range=[-25, 10]),
                yaxis=dict(range=[0, 10]),
                height=350,
                margin=dict(t=30, b=50, l=60, r=30)
            )
            st.plotly_chart(fig_scatter, use_container_width=True)
        
        # PRAVÝ PANEL - Vodíková afinita
        with col_right:
            st.subheader("💧 Vodíková afinita")
            
            # Klasifikace
            if hydrogen['color'] == "#43A047":
                st.success(f"### {hydrogen['icon']} {hydrogen['category']}")
            elif hydrogen['color'] == "#E53935":
                st.error(f"### {hydrogen['icon']} {hydrogen['category']}")
            elif hydrogen['color'] == "#FB8C00":
                st.warning(f"### {hydrogen['icon']} {hydrogen['category']}")
            else:
                st.info(f"### {hydrogen['icon']} {hydrogen['category']}")
            
            st.markdown(hydrogen['description'])
            
            st.divider()
            
            # Hodnoty
            h1, h2 = st.columns(2)
            with h1:
                st.metric("ΔH∞ [kJ/mol H]", f"{results['deltaH_inf']:.2f}")
            with h2:
                st.metric("ΔHf [kJ/mol H]", f"{results['deltaH_f']:.2f}")
            
            st.divider()
            
            # Bar chart příspěvků
            st.markdown("**Příspěvky prvků k ΔHf:**")
            contrib_data = pd.DataFrame({
                'Prvek': st.session_state.elements,
                'Příspěvek': [results['norm_comp'][el] * ELEMENTS[el]['deltaH_H'] for el in st.session_state.elements],
                'Barva': [ELEMENTS[el]['color'] for el in st.session_state.elements]
            })
            
            fig_bar = px.bar(
                contrib_data,
                x='Příspěvek',
                y='Prvek',
                orientation='h',
                color='Prvek',
                color_discrete_map={el: ELEMENTS[el]['color'] for el in st.session_state.elements}
            )
            fig_bar.update_layout(
                showlegend=False,
                height=250,
                margin=dict(t=20, b=30, l=50, r=20),
                xaxis_title="Příspěvek [kJ/mol H]"
            )
            st.plotly_chart(fig_bar, use_container_width=True)
            
            st.divider()
            
            # Export do Wordu
            st.markdown("### 📄 Export")
            
            word_buffer = create_word_report(st.session_state.elements, results, structure, hydrogen)
            st.download_button(
                label="⬇️ Stáhnout Word report",
                data=word_buffer,
                file_name=f"HEA_Report_{formula_str}_{datetime.now().strftime('%Y%m%d')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
                type="primary"
            )
    
    # Footer
    st.divider()
    st.markdown("""
    <div style="text-align: center; color: #666; font-size: 0.8rem;">
        <p>Modely: Hume-Rothery pravidla • Miedema model (binární entalpie) • Griessen-Driessen (vodíková afinita)</p>
        <p>© 2025 HEA Expert Calculator | RTI ZČU</p>
    </div>
    """, unsafe_allow_html=True)


if __name__ == "__main__":
    main()
